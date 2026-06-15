import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# CONFIG
# =========================

ALLOWED_EXTENSIONS = {
    ".py", ".js", ".ts", ".java", ".cpp", ".c", ".h",
    ".html", ".css", ".json", ".md"
}

IGNORED_DIRS = {
    ".git", "__pycache__", "node_modules",
    "dist", "build", ".idea", ".vscode", ".venv"
}

OUTPUT_FILE = "report_2cols.docx"


# =========================
# WORD HELPERS
# =========================

def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')

    for child in sectPr.findall(qn('w:cols')):
        sectPr.remove(child)

    sectPr.append(cols)


def clean_text(text):
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)


def add_code(doc, text):
    text = clean_text(text)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn('w:ascii'), "Courier New")
    run.font.size = Pt(8)


# =========================
# FILE HANDLING
# =========================

def should_include(file):
    if file.startswith("."):
        return False
    return any(file.endswith(ext) for ext in ALLOWED_EXTENSIONS)


def should_skip_dir(path):
    return any(part in IGNORED_DIRS for part in path.split(os.sep))


def read_file(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except:
        return None


# =========================
# AI DESCRIPTION
# =========================

def get_description(filename, content):
    name = filename.lower()

    # Django / backend
    if name == "manage.py":
        return "Точка входа Django проекта (запуск и управление сервером)"

    if name == "settings.py":
        return "Конфигурация проекта Django"

    if name == "urls.py":
        return "Настройка маршрутов (URL) проекта"

    if name == "views.py":
        return "Логика обработки запросов (контроллеры)"

    if name == "models.py":
        return "Модели базы данных"

    if name == "admin.py":
        return "Настройка админ-панели"

    if name == "apps.py":
        return "Конфигурация Django приложения"

    # Общие утилиты
    if "utils" in name:
        return "Вспомогательные функции проекта"

    if "config" in name:
        return "Конфигурационный файл"

    if "test" in name:
        return "Тесты проекта"

    if "main" in name:
        return "Основной запуск приложения"

    # Frontend
    if name.endswith(".js"):
        return "Фронтенд логика (JavaScript)"
    if name.endswith(".ts"):
        return "Фронтенд логика (TypeScript)"
    if name.endswith(".html"):
        return "HTML шаблон / страница"
    if name.endswith(".css"):
        return "Стили интерфейса"

    # Данные
    if name.endswith(".json"):
        return "Файл данных или конфигурации (JSON)"
    if name.endswith(".md"):
        return "Документация проекта"

    # Python общее
    if name.endswith(".py"):
        return "Python модуль"

    # По умолчанию
    return "Файл проекта"
# =========================
# MAIN
# =========================

def main():
    doc = Document()

    set_two_columns(doc.sections[0])

    doc.add_heading("Project Code Report", 0)

    for root, dirs, files in os.walk("."):
        if should_skip_dir(root):
            continue

        dirs[:] = [d for d in dirs if d not in IGNORED_DIRS]

        for file in files:
            if not should_include(file):
                continue

            path = os.path.join(root, file)
            content = read_file(path)

            if not content:
                continue

            filename = os.path.basename(path)

            # AI description
            description = get_description(filename, content)

            doc.add_heading(filename, level=2)
            doc.add_paragraph(f"Описание: {description}")
            doc.add_paragraph(f"Lines: {len(content.splitlines())}")

            add_code(doc, content)

    doc.save(OUTPUT_FILE)
    print("Готово:", OUTPUT_FILE)


if __name__ == "__main__":
    main()
